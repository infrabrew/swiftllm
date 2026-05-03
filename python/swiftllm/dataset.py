#!/usr/bin/env python3
# ==============================================================================
# PROJECT:   SWIFTLLM
# FILE:      dataset.py
# PATH:      /python/swiftllm/dataset.py
# AUTHOR:    Peter A. Aldrich Jr.
# DATE:      2026
# ------------------------------------------------------------------------------
# USED BY:
#   - python/swiftllm/__init__.py   exports DatasetIngester, IngestionConfig, etc.
#   - python/swiftllm/cli.py        `swiftllm dataset` subcommand
#   - python/swiftllm/training.py   prepare_dataset() helper, Trainer auto-ingestion
# SEE ALSO:
#   - examples/dataset_ingestion.py   end-to-end demo
# ------------------------------------------------------------------------------
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
# ==============================================================================

"""SwiftLLM Dataset Ingestion

Convert directories or individual files of any supported format into JSONL
training data ready for ``fine_tune()``, ``grpo_train()``, or ``Trainer``.

Supported input formats
-----------------------
  Plain text     : .txt  .md  .rst  .log  .text
  Code           : .py  .js  .ts  .rs  .go  .java  .c  .cpp  .cs  .rb
                   .php  .swift  .kt  .scala  .sh  .sql  .toml  .yaml  and more
  Documents      : .pdf (pdfplumber / pypdf)  .docx (python-docx)
  Web            : .html / .htm (beautifulsoup4)  .xml
  Structured     : .csv  .json  .jsonl / .ndjson

Output formats (JSONL)
----------------------
  pretraining    : {"text": "..."}
  sft_messages   : {"messages": [{"role": "system", ...}, {"role": "user", ...},
                                  {"role": "assistant", ...}]}
  sft_completion : {"prompt": "...", "completion": "..."}
  code           : {"prompt": "# <lang>\\n# File: <name>\\n\\n",
                    "completion": "<code>"}   (code files only; other files
                    fall back to pretraining)

Quick start
-----------
    from swiftllm.dataset import DatasetIngester, IngestionConfig, DatasetFormat

    result = DatasetIngester(IngestionConfig(
        input_paths=["./docs/", "./src/", "paper.pdf"],
        output_path="./data/train.jsonl",
        format=DatasetFormat.PRETRAINING,
        chunk_size=1024,
    )).ingest()

    print(f"Wrote {result.total_chunks} chunks → {result.output_path}")

Convenience function
--------------------
    from swiftllm.dataset import ingest_dataset

    result = ingest_dataset(
        input_paths="./my_project/",
        output_path="./train.jsonl",
        format="code",
    )

Optional dependencies
---------------------
    pip install pdfplumber        # PDF extraction (recommended)
    pip install pypdf             # PDF extraction (fallback)
    pip install python-docx       # DOCX extraction
    pip install beautifulsoup4    # HTML/XML extraction (falls back to regex)
"""

import csv
import hashlib
import json
import os
import re
import sys
import warnings
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Optional, Set, Tuple, Union


# ---------------------------------------------------------------------------
# Extension registry
# ---------------------------------------------------------------------------

TEXT_EXTENSIONS: frozenset = frozenset({
    ".txt", ".md", ".rst", ".log", ".text", ".readme", ".tex",
    ".asciidoc", ".adoc",
})

CODE_EXTENSIONS: frozenset = frozenset({
    ".py", ".pyi",
    ".js", ".jsx", ".mjs", ".cjs",
    ".ts", ".tsx",
    ".rs",
    ".go",
    ".java",
    ".c", ".cc", ".cpp",
    ".h", ".hpp",
    ".cs",
    ".rb",
    ".php",
    ".swift",
    ".kt", ".kts",
    ".scala",
    ".sh", ".bash", ".zsh", ".fish",
    ".lua",
    ".r",
    ".m",       # MATLAB / Objective-C
    ".sql",
    ".toml",
    ".yaml", ".yml",
    ".ini", ".cfg",
    ".dockerfile",
    ".makefile",
    ".mk",
    ".gradle",
    ".cmake",
    ".zig",
    ".nim",
    ".ex", ".exs",  # Elixir
    ".clj", ".cljs",  # Clojure
    ".hs",  # Haskell
    ".elm",
    ".dart",
    ".jl",  # Julia
    ".f90", ".f95",  # Fortran
    ".vhd", ".vhdl",  # VHDL
    ".v",   # Verilog
})

DOCUMENT_EXTENSIONS: frozenset = frozenset({".pdf", ".docx"})

WEB_EXTENSIONS: frozenset = frozenset({".html", ".htm", ".xml", ".xhtml"})

DATA_EXTENSIONS: frozenset = frozenset({".csv", ".json", ".jsonl", ".ndjson"})

ALL_EXTENSIONS: frozenset = (
    TEXT_EXTENSIONS | CODE_EXTENSIONS | DOCUMENT_EXTENSIONS
    | WEB_EXTENSIONS | DATA_EXTENSIONS
)

# Map extension → display language name (for CODE format prompts)
LANGUAGE_MAP: Dict[str, str] = {
    ".py": "python",      ".pyi": "python",
    ".js": "javascript",  ".jsx": "javascript",  ".mjs": "javascript",
    ".ts": "typescript",  ".tsx": "typescript",
    ".rs": "rust",
    ".go": "go",
    ".java": "java",
    ".c": "c",            ".cc": "cpp",           ".cpp": "cpp",
    ".h": "c",            ".hpp": "cpp",
    ".cs": "csharp",
    ".rb": "ruby",
    ".php": "php",
    ".swift": "swift",
    ".kt": "kotlin",      ".kts": "kotlin",
    ".scala": "scala",
    ".sh": "bash",        ".bash": "bash",         ".zsh": "zsh",
    ".fish": "fish",
    ".lua": "lua",
    ".r": "r",
    ".m": "matlab",
    ".sql": "sql",
    ".toml": "toml",
    ".yaml": "yaml",      ".yml": "yaml",
    ".ini": "ini",        ".cfg": "ini",
    ".dockerfile": "dockerfile",
    ".zig": "zig",
    ".nim": "nim",
    ".ex": "elixir",      ".exs": "elixir",
    ".clj": "clojure",    ".cljs": "clojure",
    ".hs": "haskell",
    ".elm": "elm",
    ".dart": "dart",
    ".jl": "julia",
    ".f90": "fortran",    ".f95": "fortran",
    ".vhd": "vhdl",       ".vhdl": "vhdl",
    ".v": "verilog",
    ".gradle": "groovy",
    ".cmake": "cmake",
}

# CSV column name synonyms used to auto-detect structured datasets
_PROMPT_COLS: Set[str] = {"prompt", "question", "instruction", "input", "query", "human"}
_COMPLETION_COLS: Set[str] = {
    "completion", "response", "answer", "output", "target",
    "label", "assistant", "gpt",
}
_MESSAGE_COLS: Set[str] = {"messages", "conversations", "dialog", "dialogue"}
_TEXT_COLS: Set[str] = {"text", "content", "document", "doc", "body", "passage"}

# Regex patterns for code block boundary detection
_CODE_BOUNDARY_RE = re.compile(
    r"^(?:def |class |fn |pub fn |func |function |const |let |var |impl "
    r"|struct |enum |interface |namespace |module |package )",
    re.MULTILINE,
)


# ---------------------------------------------------------------------------
# Public enumerations and dataclasses
# ---------------------------------------------------------------------------

class DatasetFormat(Enum):
    """JSONL output format produced by the ingester.

    PRETRAINING     ``{"text": "..."}``
    SFT_MESSAGES    ``{"messages": [...]}`` (system / user / assistant)
    SFT_COMPLETION  ``{"prompt": "...", "completion": "..."}``
    CODE            ``{"prompt": "# lang\\n# File: name\\n", "completion": code}``
                    Code files use CODE; other files fall back to PRETRAINING.
    """
    PRETRAINING = "pretraining"
    SFT_MESSAGES = "sft_messages"
    SFT_COMPLETION = "sft_completion"
    CODE = "code"


@dataclass
class IngestionConfig:
    """Configuration for a dataset ingestion run.

    Attributes
    ----------
    input_paths:
        One or more file paths or directory paths to process.  Directories
        are walked (recursively by default).
    output_path:
        Destination ``.jsonl`` file.  Created if it does not exist; parent
        directory must exist.
    format:
        Output JSONL schema.  See :class:`DatasetFormat`.
    file_extensions:
        Whitelist of extensions to process (with leading dot, e.g. ``".py"``).
        ``None`` (default) includes every extension in :data:`ALL_EXTENSIONS`.
    recursive:
        Walk directories recursively (default: ``True``).
    chunk_size:
        Maximum number of characters per JSONL record.  Long documents are
        split into overlapping chunks of this size.
    chunk_overlap:
        Character overlap between consecutive chunks to preserve context
        across boundaries.  Must be < ``chunk_size``.
    min_length:
        Chunks shorter than this are discarded (e.g. boilerplate headers).
    max_file_size_mb:
        Files larger than this limit are skipped with a warning.
    system_prompt:
        System turn content used in ``SFT_MESSAGES`` records.
    sft_user_template:
        ``SFT_MESSAGES`` / ``SFT_COMPLETION`` user-turn template.  Use
        ``{text}`` for the first half of the chunk and ``{completion}`` for
        the second half (if both are needed).  For ``SFT_COMPLETION`` the
        default splits the chunk 75 / 25 to produce a natural prompt and
        continuation.
    encoding:
        Text encoding for reading plain-text files (default: ``"utf-8"``).
    deduplicate:
        Skip chunks whose SHA-256 fingerprint has already been seen in this
        run (default: ``True``).
    strip_whitespace:
        Collapse runs of blank lines and strip leading/trailing whitespace
        (default: ``True``).
    include_metadata:
        Attach ``"_source"`` and ``"_ext"`` keys to each record (default:
        ``False``).
    verbose:
        Print per-file progress to stdout (default: ``False``).
    """
    input_paths: List[str]
    output_path: str
    format: DatasetFormat = DatasetFormat.PRETRAINING
    file_extensions: Optional[List[str]] = None
    recursive: bool = True
    chunk_size: int = 2048
    chunk_overlap: int = 128
    min_length: int = 50
    max_file_size_mb: float = 50.0
    system_prompt: str = "You are a helpful assistant."
    sft_user_template: str = "Continue the following passage:\n\n{text}"
    encoding: str = "utf-8"
    deduplicate: bool = True
    strip_whitespace: bool = True
    include_metadata: bool = False
    verbose: bool = False

    def __post_init__(self):
        if self.chunk_overlap >= self.chunk_size:
            raise ValueError(
                f"chunk_overlap ({self.chunk_overlap}) must be "
                f"less than chunk_size ({self.chunk_size})."
            )
        if isinstance(self.format, str):
            self.format = DatasetFormat(self.format)
        # Normalise extensions to lowercase with leading dot
        if self.file_extensions is not None:
            self.file_extensions = [
                ext if ext.startswith(".") else f".{ext}"
                for ext in self.file_extensions
            ]

    @property
    def _allowed_extensions(self) -> frozenset:
        if self.file_extensions is None:
            return ALL_EXTENSIONS
        return frozenset(self.file_extensions)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "input_paths": self.input_paths,
            "output_path": self.output_path,
            "format": self.format.value,
            "file_extensions": self.file_extensions,
            "recursive": self.recursive,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "min_length": self.min_length,
            "max_file_size_mb": self.max_file_size_mb,
            "system_prompt": self.system_prompt,
            "sft_user_template": self.sft_user_template,
            "encoding": self.encoding,
            "deduplicate": self.deduplicate,
            "strip_whitespace": self.strip_whitespace,
            "include_metadata": self.include_metadata,
        }

    @classmethod
    def from_dict(cls, d: Dict[str, Any]) -> "IngestionConfig":
        d = dict(d)
        if "format" in d:
            d["format"] = DatasetFormat(d["format"])
        return cls(**d)


@dataclass
class IngestionResult:
    """Statistics returned by :meth:`DatasetIngester.ingest`.

    Attributes
    ----------
    total_files_scanned:
        All files visited (including skipped).
    total_files_processed:
        Files from which at least one chunk was extracted.
    total_chunks:
        Total JSONL records written.
    total_chars:
        Total characters of raw text extracted across all files.
    skipped_files:
        ``(path, reason)`` pairs for every skipped file.
    output_path:
        Absolute path of the written ``.jsonl`` file.
    format_counts:
        Records written per file extension, e.g.
        ``{".py": 120, ".txt": 40, ".pdf": 10}``.
    """
    total_files_scanned: int = 0
    total_files_processed: int = 0
    total_chunks: int = 0
    total_chars: int = 0
    skipped_files: List[Tuple[str, str]] = field(default_factory=list)
    output_path: str = ""
    format_counts: Dict[str, int] = field(default_factory=dict)

    def summary(self) -> str:
        """Human-readable one-block summary."""
        lines = [
            f"Dataset ingestion complete",
            f"  Output          : {self.output_path}",
            f"  Files scanned   : {self.total_files_scanned}",
            f"  Files processed : {self.total_files_processed}",
            f"  Chunks written  : {self.total_chunks}",
            f"  Total chars     : {self.total_chars:,}",
        ]
        if self.format_counts:
            lines.append("  By extension    :")
            for ext, n in sorted(self.format_counts.items(), key=lambda x: -x[1]):
                lines.append(f"    {ext:<14} {n:>6} chunks")
        if self.skipped_files:
            lines.append(f"  Skipped files   : {len(self.skipped_files)}")
            for path, reason in self.skipped_files[:5]:
                lines.append(f"    {Path(path).name}: {reason}")
            if len(self.skipped_files) > 5:
                lines.append(f"    ... and {len(self.skipped_files) - 5} more")
        return "\n".join(lines)


# ---------------------------------------------------------------------------
# Core ingestion engine
# ---------------------------------------------------------------------------

class DatasetIngester:
    """Ingest files from one or more paths into JSONL training data.

    Parameters
    ----------
    config : IngestionConfig
        Full ingestion configuration.

    Example
    -------
    ::

        ingester = DatasetIngester(IngestionConfig(
            input_paths=["./data/docs/", "./notebooks/"],
            output_path="./train.jsonl",
            format=DatasetFormat.SFT_MESSAGES,
            chunk_size=1500,
            chunk_overlap=150,
        ))
        result = ingester.ingest()
        print(result.summary())
    """

    def __init__(self, config: IngestionConfig):
        self.cfg = config
        self._seen: Set[str] = set()  # SHA-256 fingerprints for dedup

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def ingest(self) -> IngestionResult:
        """Run the full ingestion pipeline and return a result summary.

        Raises
        ------
        FileNotFoundError
            If any ``input_path`` does not exist.
        OSError
            If ``output_path`` cannot be opened for writing.
        """
        result = IngestionResult(output_path=str(Path(self.cfg.output_path).resolve()))

        files = list(self._collect_files(result))

        with open(self.cfg.output_path, "w", encoding="utf-8") as out:
            for path in files:
                result.total_files_scanned += 1
                records = self._process_file(path, result)
                if records is None:
                    continue
                wrote = 0
                for rec in records:
                    out.write(json.dumps(rec, ensure_ascii=False) + "\n")
                    wrote += 1
                if wrote:
                    result.total_files_processed += 1
                    result.total_chunks += wrote
                    ext = path.suffix.lower()
                    result.format_counts[ext] = result.format_counts.get(ext, 0) + wrote

        return result

    # ------------------------------------------------------------------
    # File collection
    # ------------------------------------------------------------------

    def _collect_files(self, result: IngestionResult) -> Iterator[Path]:
        """Yield every matching file from all input_paths."""
        allowed = self.cfg._allowed_extensions
        for raw in self.cfg.input_paths:
            p = Path(raw)
            if not p.exists():
                raise FileNotFoundError(f"Input path not found: {raw!r}")
            if p.is_file():
                if p.suffix.lower() in allowed:
                    yield p
                else:
                    result.skipped_files.append((str(p), f"extension {p.suffix!r} not in allowlist"))
            elif p.is_dir():
                pattern = "**/*" if self.cfg.recursive else "*"
                for child in sorted(p.glob(pattern)):
                    if child.is_file() and child.suffix.lower() in allowed:
                        yield child

    # ------------------------------------------------------------------
    # Per-file dispatch
    # ------------------------------------------------------------------

    def _process_file(
        self, path: Path, result: IngestionResult
    ) -> Optional[List[Dict[str, Any]]]:
        """Read *path* and return a list of JSONL records, or None on skip."""
        # Size guard
        size_mb = path.stat().st_size / (1024 * 1024)
        if size_mb > self.cfg.max_file_size_mb:
            result.skipped_files.append(
                (str(path), f"file too large ({size_mb:.1f} MB > {self.cfg.max_file_size_mb} MB)")
            )
            return None

        if self.cfg.verbose:
            print(f"  Processing: {path}", flush=True)

        ext = path.suffix.lower()

        try:
            # Structured formats — may return records directly
            if ext == ".jsonl" or ext == ".ndjson":
                return self._jsonl_to_records(path, result)
            if ext == ".json":
                return self._json_to_records(path, result)
            if ext == ".csv":
                return self._csv_to_records(path, result)

            # Unstructured text — read to string, then chunk
            text = self._read_to_text(path, ext)
            if text is None:
                result.skipped_files.append((str(path), "unsupported or unreadable format"))
                return None

            result.total_chars += len(text)
            return self._text_to_records(text, path)

        except Exception as exc:  # noqa: BLE001
            result.skipped_files.append((str(path), f"read error: {exc}"))
            return None

    # ------------------------------------------------------------------
    # Text extraction (unstructured formats)
    # ------------------------------------------------------------------

    def _read_to_text(self, path: Path, ext: str) -> Optional[str]:
        """Dispatch to the correct reader; return raw text or None."""
        if ext in TEXT_EXTENSIONS:
            return self._read_plain(path)
        if ext in CODE_EXTENSIONS:
            return self._read_plain(path)  # code is plain text
        if ext == ".pdf":
            return self._read_pdf(path)
        if ext == ".docx":
            return self._read_docx(path)
        if ext in WEB_EXTENSIONS:
            return self._read_html(path)
        return None

    def _read_plain(self, path: Path) -> str:
        """Read a plain-text or code file."""
        return path.read_text(encoding=self.cfg.encoding, errors="replace")

    def _read_pdf(self, path: Path) -> str:
        """Extract text from a PDF, trying pdfplumber then pypdf then PyPDF2."""
        # Attempt 1: pdfplumber (best layout fidelity)
        try:
            import pdfplumber  # type: ignore
            with pdfplumber.open(str(path)) as pdf:
                pages = []
                for page in pdf.pages:
                    txt = page.extract_text()
                    if txt:
                        pages.append(txt)
            return "\n\n".join(pages)
        except ImportError:
            pass
        except Exception as exc:
            warnings.warn(f"pdfplumber failed on {path.name}: {exc}; trying pypdf")

        # Attempt 2: pypdf
        try:
            import pypdf  # type: ignore
            reader = pypdf.PdfReader(str(path))
            return "\n\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        except ImportError:
            pass
        except Exception as exc:
            warnings.warn(f"pypdf failed on {path.name}: {exc}; trying PyPDF2")

        # Attempt 3: PyPDF2 (legacy)
        try:
            import PyPDF2  # type: ignore
            with open(path, "rb") as fh:
                reader = PyPDF2.PdfReader(fh)
                return "\n\n".join(
                    page.extract_text() or "" for page in reader.pages
                )
        except ImportError:
            raise ImportError(
                f"PDF support requires a PDF library.  Install one with:\n"
                f"  pip install pdfplumber   # recommended\n"
                f"  pip install pypdf\n"
                f"  pip install PyPDF2       # legacy"
            )

    def _read_docx(self, path: Path) -> str:
        """Extract text from a .docx file using python-docx."""
        try:
            import docx  # type: ignore
            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract table cell text
            for table in doc.tables:
                for row in table.rows:
                    cell_texts = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cell_texts:
                        paragraphs.append(" | ".join(cell_texts))
            return "\n\n".join(paragraphs)
        except ImportError:
            raise ImportError(
                "DOCX support requires python-docx.  Install with:\n"
                "  pip install python-docx"
            )

    def _read_html(self, path: Path) -> str:
        """Extract visible text from HTML/XML.  Uses BeautifulSoup if available."""
        raw = path.read_text(encoding=self.cfg.encoding, errors="replace")
        try:
            from bs4 import BeautifulSoup  # type: ignore
            soup = BeautifulSoup(raw, "html.parser")
            # Remove non-content tags
            for tag in soup(["script", "style", "meta", "head",
                              "nav", "footer", "header", "noscript"]):
                tag.decompose()
            return soup.get_text(separator="\n", strip=True)
        except ImportError:
            # Fallback: strip HTML tags with regex
            text = re.sub(r"<[^>]+>", " ", raw)
            text = re.sub(r"&[a-zA-Z]+;", " ", text)  # HTML entities
            return re.sub(r"\s{2,}", " ", text).strip()

    # ------------------------------------------------------------------
    # Structured format readers (return records directly)
    # ------------------------------------------------------------------

    def _jsonl_to_records(
        self, path: Path, result: IngestionResult
    ) -> List[Dict[str, Any]]:
        """Pass-through or remap existing JSONL to the configured output format."""
        records: List[Dict[str, Any]] = []
        with open(path, encoding=self.cfg.encoding, errors="replace") as fh:
            for lineno, line in enumerate(fh, 1):
                line = line.strip()
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                except json.JSONDecodeError:
                    continue
                if not isinstance(obj, dict):
                    continue

                remapped = self._remap_record(obj, str(path))
                if remapped is not None:
                    records.extend(remapped)
        result.total_chars += path.stat().st_size
        return records

    def _json_to_records(
        self, path: Path, result: IngestionResult
    ) -> List[Dict[str, Any]]:
        """Handle a .json file: list-of-records or single object."""
        text = path.read_text(encoding=self.cfg.encoding, errors="replace")
        try:
            data = json.loads(text)
        except json.JSONDecodeError as exc:
            raise ValueError(f"Invalid JSON in {path.name}: {exc}") from exc

        result.total_chars += len(text)
        items = data if isinstance(data, list) else [data]
        records: List[Dict[str, Any]] = []
        for obj in items:
            if isinstance(obj, dict):
                remapped = self._remap_record(obj, str(path))
                if remapped:
                    records.extend(remapped)
        return records

    def _csv_to_records(
        self, path: Path, result: IngestionResult
    ) -> List[Dict[str, Any]]:
        """Convert a CSV into JSONL records.

        Detects structured columns (prompt/completion, messages, text) and
        remaps them.  Falls back to concatenating all values as plain text.
        """
        records: List[Dict[str, Any]] = []
        with open(path, encoding=self.cfg.encoding, errors="replace", newline="") as fh:
            reader = csv.DictReader(fh)
            if reader.fieldnames is None:
                return records

            col_lower = {c.lower().strip(): c for c in reader.fieldnames if c}
            prompt_col = next((col_lower[k] for k in _PROMPT_COLS if k in col_lower), None)
            compl_col  = next((col_lower[k] for k in _COMPLETION_COLS if k in col_lower), None)
            msg_col    = next((col_lower[k] for k in _MESSAGE_COLS if k in col_lower), None)
            text_col   = next((col_lower[k] for k in _TEXT_COLS if k in col_lower), None)

            for row in reader:
                obj: Dict[str, Any] = {}

                if msg_col and row.get(msg_col):
                    # messages column: may be JSON string
                    raw_msg = row[msg_col]
                    try:
                        obj["messages"] = json.loads(raw_msg)
                    except (json.JSONDecodeError, TypeError):
                        obj["messages"] = raw_msg  # leave as-is; remap handles it
                elif prompt_col and compl_col:
                    obj["prompt"] = row.get(prompt_col, "").strip()
                    obj["completion"] = row.get(compl_col, "").strip()
                elif text_col:
                    obj["text"] = row.get(text_col, "").strip()
                else:
                    # Concatenate all column values as free text
                    parts = [v.strip() for v in row.values() if v and v.strip()]
                    if parts:
                        obj["text"] = "  ".join(parts)

                if obj:
                    remapped = self._remap_record(obj, str(path))
                    if remapped:
                        records.extend(remapped)

        result.total_chars += path.stat().st_size
        return records

    # ------------------------------------------------------------------
    # Record mapping (structured → output format)
    # ------------------------------------------------------------------

    def _remap_record(
        self, obj: Dict[str, Any], source: str
    ) -> Optional[List[Dict[str, Any]]]:
        """Convert an already-parsed dict to the configured output format.

        Handles objects that already have ``text``, ``prompt``/``completion``,
        or ``messages`` keys, remapping to the target format.
        """
        fmt = self.cfg.format

        # Detect what keys are present
        has_text       = "text" in obj and obj["text"]
        has_pc         = "prompt" in obj and "completion" in obj
        has_messages   = "messages" in obj and obj["messages"]

        # Normalise input to (prompt_str, completion_str) or text_str
        if has_messages:
            msgs = obj["messages"]
            if isinstance(msgs, list):
                # Already in messages format → keep or convert
                if fmt == DatasetFormat.SFT_MESSAGES:
                    return [self._maybe_meta({"messages": msgs}, source, "jsonl")]
                # Convert to text by concatenating role: content
                text = "\n".join(
                    f"{m.get('role', 'user').capitalize()}: {m.get('content', '')}"
                    for m in msgs if isinstance(m, dict)
                )
                has_text, obj["text"] = True, text

        if has_pc:
            prompt = str(obj["prompt"]).strip()
            completion = str(obj["completion"]).strip()
            if not prompt and not completion:
                return None
            return self._pc_to_records(prompt, completion, source, "jsonl")

        if has_text:
            text = str(obj["text"]).strip()
            if len(text) < self.cfg.min_length:
                return None
            return self._text_to_records(text, Path(source), ext_override="jsonl")

        # Unknown structure — try to stringify all values
        parts = []
        for v in obj.values():
            if isinstance(v, str) and v.strip():
                parts.append(v.strip())
        if parts:
            text = " ".join(parts)
            if len(text) >= self.cfg.min_length:
                return self._text_to_records(text, Path(source), ext_override="jsonl")
        return None

    # ------------------------------------------------------------------
    # Text → records
    # ------------------------------------------------------------------

    def _text_to_records(
        self,
        text: str,
        path: Path,
        ext_override: Optional[str] = None,
    ) -> List[Dict[str, Any]]:
        """Chunk *text* and convert each chunk to the configured output format."""
        if self.cfg.strip_whitespace:
            text = self._normalise_whitespace(text)

        ext = ext_override or path.suffix.lower()
        is_code = ext in CODE_EXTENSIONS
        lang = LANGUAGE_MAP.get(ext, "")

        chunks = self._split_code(text, lang) if is_code else self._chunk_text(text)

        records: List[Dict[str, Any]] = []
        for chunk in chunks:
            chunk = chunk.strip()
            if len(chunk) < self.cfg.min_length:
                continue
            if self.cfg.deduplicate and not self._is_unique(chunk):
                continue

            rec = self._chunk_to_record(chunk, path, ext, lang, is_code)
            if rec is not None:
                records.append(rec)

        return records

    def _pc_to_records(
        self,
        prompt: str,
        completion: str,
        source: str,
        ext: str,
    ) -> List[Dict[str, Any]]:
        """Convert an explicit prompt/completion pair to the output format."""
        fmt = self.cfg.format
        key = hashlib.sha256((prompt + completion).encode()).hexdigest()
        if self.cfg.deduplicate and key in self._seen:
            return []
        self._seen.add(key)

        if fmt == DatasetFormat.PRETRAINING:
            text = f"{prompt}\n{completion}"
            return [self._maybe_meta({"text": text}, source, ext)]
        if fmt == DatasetFormat.SFT_COMPLETION:
            return [self._maybe_meta({"prompt": prompt, "completion": completion}, source, ext)]
        if fmt in (DatasetFormat.SFT_MESSAGES, DatasetFormat.CODE):
            return [self._maybe_meta({
                "messages": [
                    {"role": "system",    "content": self.cfg.system_prompt},
                    {"role": "user",      "content": prompt},
                    {"role": "assistant", "content": completion},
                ]
            }, source, ext)]
        return []

    def _chunk_to_record(
        self,
        chunk: str,
        path: Path,
        ext: str,
        lang: str,
        is_code: bool,
    ) -> Optional[Dict[str, Any]]:
        """Convert a single text chunk to the configured output format record."""
        fmt = self.cfg.format

        if fmt == DatasetFormat.PRETRAINING:
            return self._maybe_meta({"text": chunk}, str(path), ext)

        if fmt == DatasetFormat.CODE:
            if is_code:
                header = f"# {lang}\n# File: {path.name}\n\n" if lang else f"# File: {path.name}\n\n"
                return self._maybe_meta(
                    {"prompt": header, "completion": chunk}, str(path), ext
                )
            # Non-code files fall back to pretraining under CODE format
            return self._maybe_meta({"text": chunk}, str(path), ext)

        if fmt == DatasetFormat.SFT_COMPLETION:
            # Split chunk ~75/25 for natural prompt / completion split
            split = max(self.cfg.min_length, int(len(chunk) * 0.75))
            # Prefer splitting at a sentence boundary near the split point
            split = self._find_sentence_boundary(chunk, split)
            prompt_text = self.cfg.sft_user_template.format(text=chunk[:split])
            completion_text = chunk[split:].strip()
            if not completion_text:
                # Chunk too short to split — treat as pretraining
                return self._maybe_meta({"text": chunk}, str(path), ext)
            return self._maybe_meta(
                {"prompt": prompt_text, "completion": completion_text}, str(path), ext
            )

        if fmt == DatasetFormat.SFT_MESSAGES:
            split = max(self.cfg.min_length, int(len(chunk) * 0.75))
            split = self._find_sentence_boundary(chunk, split)
            user_content = self.cfg.sft_user_template.format(text=chunk[:split])
            assistant_content = chunk[split:].strip()
            if not assistant_content:
                assistant_content = chunk  # use full chunk as assistant turn
                user_content = self.cfg.sft_user_template.format(text="")
            return self._maybe_meta({
                "messages": [
                    {"role": "system",    "content": self.cfg.system_prompt},
                    {"role": "user",      "content": user_content},
                    {"role": "assistant", "content": assistant_content},
                ]
            }, str(path), ext)

        return None

    # ------------------------------------------------------------------
    # Chunking helpers
    # ------------------------------------------------------------------

    def _chunk_text(self, text: str) -> List[str]:
        """Split *text* into overlapping character-level chunks."""
        size = self.cfg.chunk_size
        overlap = self.cfg.chunk_overlap
        stride = size - overlap
        if len(text) <= size:
            return [text]
        chunks = []
        start = 0
        while start < len(text):
            end = start + size
            chunk = text[start:end]
            # Try to end on a sentence boundary to avoid cutting mid-word
            if end < len(text):
                boundary = self._find_sentence_boundary(chunk, len(chunk))
                chunk = chunk[:boundary] if boundary > overlap else chunk
            chunks.append(chunk)
            start += stride
        return chunks

    def _split_code(self, text: str, lang: str) -> List[str]:
        """Split code at top-level definition boundaries when possible."""
        size = self.cfg.chunk_size
        if len(text) <= size:
            return [text]

        # Find positions of top-level definitions
        boundaries = [0] + [m.start() for m in _CODE_BOUNDARY_RE.finditer(text)]

        chunks: List[str] = []
        current_start = 0

        for boundary in boundaries[1:]:
            segment = text[current_start:boundary]
            if len(segment) >= size:
                # Segment alone is too large — fall back to character chunking
                chunks.extend(self._chunk_text(segment))
                current_start = boundary
            elif len(text[current_start:boundary + size]) > size:
                # Adding next definition would overflow — flush current
                chunks.append(segment)
                current_start = boundary

        # Flush remainder
        remainder = text[current_start:]
        if remainder.strip():
            if len(remainder) > size:
                chunks.extend(self._chunk_text(remainder))
            else:
                chunks.append(remainder)

        return chunks if chunks else [text]

    @staticmethod
    def _find_sentence_boundary(text: str, target: int) -> int:
        """Return an index near *target* that falls on a sentence end."""
        # Search backwards up to 15% of target for a sentence terminator
        window = max(1, target // 7)
        for i in range(target, max(target - window, 0), -1):
            if i < len(text) and text[i] in ".!?\n":
                return i + 1
        return target

    # ------------------------------------------------------------------
    # Whitespace normalisation & deduplication
    # ------------------------------------------------------------------

    @staticmethod
    def _normalise_whitespace(text: str) -> str:
        """Collapse >2 consecutive blank lines; strip leading/trailing space."""
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _is_unique(self, chunk: str) -> bool:
        """Return True if chunk has not been seen; register its fingerprint."""
        key = hashlib.sha256(chunk.encode()).hexdigest()
        if key in self._seen:
            return False
        self._seen.add(key)
        return True

    # ------------------------------------------------------------------
    # Metadata attachment
    # ------------------------------------------------------------------

    def _maybe_meta(
        self, record: Dict[str, Any], source: str, ext: str
    ) -> Dict[str, Any]:
        """Optionally attach _source / _ext fields to a record."""
        if self.cfg.include_metadata:
            record["_source"] = Path(source).name
            record["_ext"] = ext
        return record


# ---------------------------------------------------------------------------
# Convenience function
# ---------------------------------------------------------------------------

def ingest_dataset(
    input_paths: Union[str, List[str]],
    output_path: str,
    format: Union[str, DatasetFormat] = DatasetFormat.PRETRAINING,
    chunk_size: int = 2048,
    chunk_overlap: int = 128,
    min_length: int = 50,
    recursive: bool = True,
    file_extensions: Optional[List[str]] = None,
    system_prompt: str = "You are a helpful assistant.",
    sft_user_template: str = "Continue the following passage:\n\n{text}",
    deduplicate: bool = True,
    include_metadata: bool = False,
    verbose: bool = False,
    **kwargs: Any,
) -> IngestionResult:
    """Ingest files into a JSONL training dataset.

    A thin wrapper around :class:`DatasetIngester` for common one-liner use.

    Parameters
    ----------
    input_paths:
        Single path string or list of paths (files or directories).
    output_path:
        Destination ``.jsonl`` file.
    format:
        Output format: ``"pretraining"``, ``"sft_messages"``,
        ``"sft_completion"``, or ``"code"``.  Accepts string or
        :class:`DatasetFormat` enum.
    chunk_size:
        Max characters per chunk (default: 2048).
    chunk_overlap:
        Character overlap between chunks (default: 128).
    min_length:
        Minimum chunk length; shorter chunks are discarded (default: 50).
    recursive:
        Walk directories recursively (default: ``True``).
    file_extensions:
        Whitelist of extensions, e.g. ``[".py", ".md"]``.
        ``None`` = all supported.
    system_prompt:
        System turn for ``sft_messages`` format.
    sft_user_template:
        User-turn template; ``{text}`` is replaced with the prompt portion.
    deduplicate:
        Skip duplicate chunks (default: ``True``).
    include_metadata:
        Add ``_source`` / ``_ext`` keys to each record.
    verbose:
        Print per-file progress.
    **kwargs:
        Forwarded to :class:`IngestionConfig`.

    Returns
    -------
    IngestionResult
        Statistics including chunk count and skipped files.

    Examples
    --------
    Pretraining from a docs directory::

        result = ingest_dataset("./docs/", "./train.jsonl")
        print(result.summary())

    Code fine-tuning from a source tree::

        result = ingest_dataset(
            input_paths=["./src/", "./tests/"],
            output_path="./code_train.jsonl",
            format="code",
            file_extensions=[".py", ".rs"],
        )

    SFT from mixed sources::

        result = ingest_dataset(
            input_paths=["papers.pdf", "qa_pairs.csv", "./notes/"],
            output_path="./sft_train.jsonl",
            format="sft_completion",
            chunk_size=1024,
        )
    """
    if isinstance(input_paths, str):
        input_paths = [input_paths]

    cfg = IngestionConfig(
        input_paths=input_paths,
        output_path=output_path,
        format=DatasetFormat(format) if isinstance(format, str) else format,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        min_length=min_length,
        recursive=recursive,
        file_extensions=file_extensions,
        system_prompt=system_prompt,
        sft_user_template=sft_user_template,
        deduplicate=deduplicate,
        include_metadata=include_metadata,
        verbose=verbose,
        **kwargs,
    )
    return DatasetIngester(cfg).ingest()


# ---------------------------------------------------------------------------
# CLI entry point (called by cli.py)
# ---------------------------------------------------------------------------

def _cli_ingest(args) -> None:  # noqa: ANN001
    """Handler for ``swiftllm dataset`` called from cli.py."""
    input_paths = args.input if isinstance(args.input, list) else [args.input]
    fmt = DatasetFormat(args.format)

    ext_list = None
    if getattr(args, "extensions", None):
        ext_list = [e.strip() for e in args.extensions.split(",")]

    cfg = IngestionConfig(
        input_paths=input_paths,
        output_path=args.output,
        format=fmt,
        file_extensions=ext_list,
        recursive=not getattr(args, "no_recursive", False),
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap,
        min_length=args.min_length,
        max_file_size_mb=args.max_file_size_mb,
        system_prompt=getattr(args, "system_prompt", "You are a helpful assistant."),
        deduplicate=not getattr(args, "no_dedup", False),
        include_metadata=getattr(args, "include_metadata", False),
        verbose=getattr(args, "verbose", False),
    )

    print(f"SwiftLLM Dataset Ingester")
    print(f"  Input paths : {', '.join(cfg.input_paths)}")
    print(f"  Output      : {cfg.output_path}")
    print(f"  Format      : {cfg.format.value}")
    print(f"  Chunk size  : {cfg.chunk_size} chars  (overlap: {cfg.chunk_overlap})")
    print()

    ingester = DatasetIngester(cfg)
    result = ingester.ingest()
    print()
    print(result.summary())


# ------------------------------------------------------------------------------
# END OF FILE: dataset.py
# REPO PATH:   /swiftllm/python/swiftllm/dataset.py
# SEE ALSO:
#   examples/dataset_ingestion.py     end-to-end demo
#   python/swiftllm/training.py       prepare_dataset() helper, Trainer integration
#   python/swiftllm/cli.py            `swiftllm dataset` subcommand
# (c) 2026 SWIFTLLM | Apache 2.0 License
# ------------------------------------------------------------------------------
